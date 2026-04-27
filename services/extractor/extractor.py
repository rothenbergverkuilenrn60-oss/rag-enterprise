# =============================================================================
# services/extractor/extractor.py  (企业级重写版)
# STAGE 2 — 企业级文档提取（500强标准）
#
# 旧版问题：PDF → PyMuPDF → 纯文本（直接解析，有4个致命缺陷）
# 新版方案：PDF → 类型判断(数字/扫描件) → 对应处理链路
#              数字PDF: 页眉页脚过滤 + 双栏重排 + pdfplumber表格还原
#              扫描件:  PaddleOCR/Tesseract + 阅读顺序重排
# =============================================================================
from __future__ import annotations
import asyncio
import json
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal
from loguru import logger
from tenacity import retry, stop_after_attempt, wait_random_exponential
from torch import device

from config.settings import settings
from utils.models import RawDocument, ExtractedContent, DocType
from utils.logger import log_latency
from services.preprocessor.cleaner import clean_text
from services.extractor.image_extractor import extract_images_from_pdf, get_image_extractor


# ══════════════════════════════════════════════════════════════════════════════
# STEP 0: 判断 PDF 是否为扫描件
# ══════════════════════════════════════════════════════════════════════════════
def is_scanned_pdf(file_path: Path, sample_pages: int = 3) -> bool:
    """
    抽样前N页，计算 "有效字符数 / 页面面积" 密度。
    密度 < 0.01 字符/pt² 则认定为扫描件。
    """
    try:
        import fitz
        doc = fitz.open(str(file_path))
        total_chars = 0
        total_area = 0.0
        pages_checked = min(sample_pages, doc.page_count)
        for i in range(pages_checked):
            page = doc[i]
            total_chars += len(page.get_text("text").strip())   # 获取页面纯文本并计算字符数
            rect = page.rect
            total_area += rect.width * rect.height
        doc.close()     # 关闭文档，释放资源
        if total_area == 0:
            return True
        density = total_chars / total_area
        logger.debug(f"PDF char density={density:.6f}")
        if density < 0.01:
            logger.info(f"[PDF判断] 扫描件: density={density:.6f} file={file_path.name}")
            return True
        return False
    except Exception as exc:
        logger.warning(f"is_scanned_pdf check failed: {exc}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1A: 数字 PDF — 结构化提取
# ══════════════════════════════════════════════════════════════════════════════
def _detect_header_footer_texts(file_path: Path, max_pages: int = 10) -> set:
    """1
    检测页眉/页脚噪声：在多页顶部/底部区域重复出现的文本行。
    出现次数 >= 总页数*40% 则认定为噪声。
    """
    try:
        import fitz
        doc = fitz.open(str(file_path))
        pages_to_scan = min(max_pages, doc.page_count)      # 最多扫描 max_pages 页
        line_counter: Counter = Counter()                   # 统计每行文本出现次数
        for i in range(pages_to_scan):
            page = doc[i]
            page_h = page.rect.height
            top_zone = fitz.Rect(0, 0, page.rect.width, page_h * 0.08)
            bot_zone = fitz.Rect(0, page_h * 0.92, page.rect.width, page_h)
            for zone in [top_zone, bot_zone]:
                for line in page.get_text("text", clip=zone).strip().split("\n"):   # clip: 仅提取指定区域文本,line: 每行文本
                    line = line.strip()
                    if len(line) > 3:
                        line_counter[line] += 1     # 统计每行文本出现次数
        doc.close()
        threshold = pages_to_scan * 0.4
        noise = {line for line, cnt in line_counter.items() if cnt >= threshold}   # 过滤出现次数 >= 总页数*40% 的文本行
        if noise:
            logger.debug(f"检测到 {len(noise)} 条页眉/页脚噪声")
        return noise
    except Exception as exc:
        logger.warning(f"页眉页脚检测失败: {exc}")
        return set()


def _is_multi_column(blocks: list, page_width: float) -> bool:      # 判断是否为多栏文本,block: 文字块坐标
    if len(blocks) < 4:     # 至少需要4个文字块才能判断是否为多栏文本
        return False
    mid = page_width / 2    # 页面宽度中心
    left = [b for b in blocks if b["x1"] < mid * 1.1]   # 左侧文字块
    right = [b for b in blocks if b["x0"] > mid * 0.9]   # 右侧文字块
    return min(len(left), len(right)) / len(blocks) >= 0.3


def _sort_blocks_multi_column(blocks: list, page_width: float) -> list:
    mid = page_width / 2
    left_col = sorted([b for b in blocks if (b["x0"]+b["x1"])/2 < mid], key=lambda b: b["y0"])  # 左侧文字块
    right_col = sorted([b for b in blocks if (b["x0"]+b["x1"])/2 >= mid], key=lambda b: b["y0"])
    return left_col + right_col


def _extract_pdf_digital(file_path: Path) -> dict:
    """
    数字PDF完整链路:
    1. 检测页眉页脚噪声
    2. PyMuPDF逐页提取文字块(带坐标)
    3. 过滤噪声 + 多栏重排
    4. pdfplumber 表格结构还原
    5. 组装干净正文
    """
    import fitz
    import pdfplumber

    noise_texts = _detect_header_footer_texts(file_path)
    doc = fitz.open(str(file_path))
    metadata = doc.metadata or {}
    all_pages_text: list = []

    for page_idx in range(doc.page_count):
        page = doc[page_idx]
        page_width = page.rect.width
        page_height = page.rect.height
        raw_blocks = page.get_text("blocks")   # 获取页面所有文本块，每个块包含 (x0,y0,x1,y1,text,block_no,block_type)
        text_blocks = [
            {"x0": b[0], "y0": b[1], "x1": b[2], "y1": b[3], "text": b[4].strip()}
            for b in raw_blocks
            if b[6] == 0 and b[4].strip()  # b[6]==0 表示文本块（0=文本，1=图片），过滤空块
        ]
        # 位置过滤(页眉页脚区域) + 文本噪声过滤
        header_bottom = page_height * 0.08
        footer_top = page_height * 0.92
        text_blocks = [
            b for b in text_blocks
            if not (b["y1"] < header_bottom or b["y0"] > footer_top)
            and not any(noise in b["text"] for noise in noise_texts)    # 不能有页眉页脚噪声文本
        ]
        if not text_blocks:
            continue
        if _is_multi_column(text_blocks, page_width):
            text_blocks = _sort_blocks_multi_column(text_blocks, page_width)
        else:
            text_blocks = sorted(text_blocks, key=lambda b: (b["y0"], b["x0"]))  # 按 y0, x0 排序
        page_text = "\n".join(b["text"] for b in text_blocks)
        if page_text.strip():
            all_pages_text.append(f"[第{page_idx+1}页]\n{page_text}")

    doc.close()

    # pdfplumber 表格
    all_tables: list = []
    table_text_parts: list = []                     # 表格文本块
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            for i, page in enumerate(pdf.pages):    # i: 页索引,page: 页面对象
                tables = page.extract_tables()
                for tbl in tables:                  # tbl: 表格对象,tables: 页面所有表格
                    if not tbl or not any(any(c for c in row) for row in tbl):  # 只要不是有非空单元格
                        continue
                    clean_rows = [[str(c or "").strip() for c in row] for row in tbl]   # 清理表格数据，转换为字符串并去空格
                    all_tables.append({"page": i+1, "rows": clean_rows})
                    # 外层的 \n 是在拼接“表格的行”，内层的 " | " 是在拼接“行里的单元格”
                    rows_str = "\n".join(" | ".join(r) for r in clean_rows if any(c.strip() for c in r))  # 把表格每行用 '|' 分隔拼成文本，空行跳过
                    if rows_str.strip():    # 过滤空表格
                        table_text_parts.append(f"[第{i+1}页·表格]\n{rows_str}")
    except Exception as exc:
        logger.warning(f"pdfplumber表格提取失败: {exc}")

    body_text = "\n\n".join(all_pages_text)
    if table_text_parts:
        body_text += "\n\n" + "\n\n".join(table_text_parts)   # 表格文本块之间用2空行隔开

    return {
        "body_text": body_text,
        "tables": all_tables,
        "pages": len(all_pages_text),
        "title": metadata.get("title", file_path.stem),
        "author": metadata.get("author", ""),
        "created_date": metadata.get("creationDate", ""),
        "metadata": metadata,
        "engine": "pymupdf+pdfplumber(digital)",
        "noise_texts_filtered": len(noise_texts),
        "tables_extracted": len(all_tables),
    }


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1B: 扫描件 PDF — OCR + 版面分析
# ══════════════════════════════════════════════════════════════════════════════
def _extract_pdf_scanned_paddleocr(file_path: Path) -> dict:
    """
    PaddleOCR处理扫描件:
    1. PDF → 逐页渲染为图片(150 DPI)
    2. PaddleOCR 识别(中英混合)
    3. 按y坐标排序(版面重排)
    4. 过滤低置信度结果(< 0.6)
    """
    try:
        import fitz
        from paddleocr import PaddleOCR
        import numpy as np

        # 初始化 PaddleOCR：use_angle_cls 自动矫正倾斜文字，lang='ch' 中英混合识别，show_log=False 关闭日志
        ocr = PaddleOCR(use_angle_cls=True, lang="ch", device="cpu",show_log=False)
        doc = fitz.open(str(file_path))
        all_pages_text: list = []

        for page_idx in range(doc.page_count):
            page = doc[page_idx]
            mat = fitz.Matrix(150/72, 150/72)   # x方向150 DPI, y方向150 DPI
            pix = page.get_pixmap(matrix=mat, alpha=False)  # 把 PDF 页面渲染为图片像素数组，alpha=False 不要透明通道
            # pix.sample:一段连续的字节流，包含每个像素的 RGB 值
            # np.frombuffer:将字节流转换为 numpy 数组,而不复制内存，dtype=uint8 表示每个字节表示一个 8 位无符号整数，图像的每个通道都是 0–255，所以用 uint8。
            img_array = np.frombuffer(
                pix.samples, dtype=np.uint8
            ).reshape(pix.height, pix.width, 3) # 转换为 numpy 数组，形状为 (高度, 宽度, 通道数)

            result = ocr.predict(img_array, cls=True)   # 推理图片中的文本
            if not result or not result[0]:
                continue

            blocks = []
            for line in result[0]:      # result[0]: 页面所有文本行,result[1]: 所有表格
                if line is None:
                    continue
                bbox, (text, conf) = line       # bbox: 文本框坐标,text: 文本内容,conf: 置信度
                if conf < 0.6 or not text.strip():
                    continue
                '''
                bbox = [
                    (x0, y0),   # 左上角
                    (x1, y1),   # 右上角
                    (x2, y2),   # 右下角
                    (x3, y3)    # 左下角
                ]
                '''
                y_center = (bbox[0][1] + bbox[2][1]) / 2
                x_start = bbox[0][0]
                blocks.append({"text": text.strip(), "y": y_center, "x": x_start})

            # y坐标/20目的是把“同一行的文本块”归为同一组,四舍五入后得到一个“行号”,b["x"]是在同一行内，按 x 坐标排序
            blocks.sort(key=lambda b: (round(b["y"] / 20), b["x"]))
            page_text = "\n".join(b["text"] for b in blocks)        # 按行拼接文本块
            if page_text.strip():
                all_pages_text.append(f"[第{page_idx+1}页·OCR]\n{page_text}")

        doc.close()
        return {
            "body_text": "\n\n".join(all_pages_text),
            "tables": [],
            "pages": len(all_pages_text),
            "title": file_path.stem,
            "engine": "paddleocr(scanned)",
        }

    except ImportError:
        logger.warning("PaddleOCR未安装，回退Tesseract")
        return _extract_pdf_scanned_tesseract(file_path)


def _extract_pdf_scanned_tesseract(file_path: Path) -> dict:
    """Tesseract OCR 回退方案（轻量，无需GPU）。"""
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io

        doc = fitz.open(str(file_path))
        all_pages_text: list = []
        for page_idx in range(doc.page_count):
            page = doc[page_idx]
            mat = fitz.Matrix(150/72, 150/72)
            '''
            PDF Page
            ↓  page.get_pixmap()
            Pixmap (pix)
            ↓  pix.tobytes("png")
            PNG bytes
            ↓  io.BytesIO(...)
            Memory file object
            ↓  Image.open(...)
            PIL Image
            '''
            pix = page.get_pixmap(matrix=mat, alpha=False)
            # 把 PDF 页面渲染成一张 PNG 图片,BytesIO 是一个 内存中的文件对象,把字节流包装成“文件一样的对象让 Pillow（PIL）可以像读文件一样读取它
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(
                img, lang=settings.extractor_ocr_lang, config="--psm 3"
            )
            if text.strip():
                all_pages_text.append(f"[第{page_idx+1}页·Tesseract]\n{text.strip()}")
        doc.close()
        return {
            "body_text": "\n\n".join(all_pages_text),
            "tables": [],
            "pages": len(all_pages_text),
            "title": file_path.stem,
            "engine": "tesseract(scanned)",
        }
    except ImportError:
        return {
            "body_text": "",
            "tables": [],
            "pages": 0,
            "title": file_path.stem,
            "engine": "none",
            "extraction_errors": ["OCR引擎未安装: pip install paddlepaddle paddleocr 或 pytesseract pillow"],
        }


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1C-Vision: Claude Vision 扫描件提取
# ══════════════════════════════════════════════════════════════════════════════
_VISION_OCR_SYSTEM = """\
你是专业的文档文字提取助手。
请将图片中的所有文字内容完整提取出来：
  - 保持段落结构，段落间用空行隔开
  - 表格用 Markdown 格式输出（| 列1 | 列2 |）
  - 标题用 # 标记
  - 只输出提取的文字，不要添加任何说明或注释
"""

async def _extract_pdf_vision_async(file_path: Path, llm_client) -> dict:
    """
    Claude Vision 扫描件提取。

    相比 OCR（PaddleOCR/Tesseract）的优势：
      - 原生理解表格布局，不需要单独的表格解析步骤
      - 理解多栏版式、不规则排版
      - 识别手写体、特殊字体
      - 理解图表中的文字说明
      - 无需额外依赖（不需要 paddlepaddle/tesseract）

    注意：每页渲染为 PNG 后通过 Base64 传给 Claude Vision API。
    """
    import fitz
    import base64

    doc = fitz.open(str(file_path))
    all_pages_text: list[str] = []
    max_pages = min(doc.page_count, 50)   # 单次最多处理 50 页，避免超时

    for page_idx in range(max_pages):
        page = doc[page_idx]
        mat  = fitz.Matrix(150 / 72, 150 / 72)          # 150 DPI，清晰度与速度平衡
        pix  = page.get_pixmap(matrix=mat, alpha=False)
        img_b64 = base64.b64encode(pix.tobytes("png")).decode()

        try:
            text = await llm_client.chat_with_vision(
                system=_VISION_OCR_SYSTEM,
                image_b64=img_b64,
                query="请提取这一页中的全部文字内容。",
                task_type="generate",
            )
            if text.strip():
                all_pages_text.append(f"[第{page_idx + 1}页·Vision]\n{text.strip()}")
        except Exception as exc:
            logger.warning(f"[Vision] page {page_idx + 1} failed: {exc}")

    doc.close()
    return {
        "body_text": "\n\n".join(all_pages_text),
        "tables":    [],
        "pages":     len(all_pages_text),
        "title":     file_path.stem,
        "engine":    "claude-vision(scanned)",
    }


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1C: 企业级 PDF 统一调度入口
# ══════════════════════════════════════════════════════════════════════════════
def _extract_pdf_enterprise(file_path: Path) -> dict:
    """
    PDF 提取统一调度入口。

    自动检测文档类型并路由到对应处理链路：
      数字 PDF → PyMuPDF（文字块+坐标）+ pdfplumber（表格结构）双引擎
      扫描件   → OCR 引擎（由 settings.ocr_engine 决定具体引擎）

    ocr_engine 可选值：
      "auto"      先尝试 PaddleOCR，未安装自动降级到 Tesseract
      "paddle"    强制 PaddleOCR
      "tesseract" 强制 Tesseract
      "none"      跳过扫描件（返回空文本，适合确认文档全为数字 PDF 的场景）
    """
    scanned = is_scanned_pdf(file_path)

    if not scanned:
        logger.info(f"[PDF] 数字PDF → 双引擎提取: {file_path.name}")
        return _extract_pdf_digital(file_path)

    # 扫描件路由
    ocr_engine = getattr(settings, "ocr_engine", "auto")
    logger.info(f"[PDF] 扫描件 → OCR engine={ocr_engine}: {file_path.name}")

    if ocr_engine == "none":
        logger.warning(f"[PDF] ocr_engine=none，跳过扫描件: {file_path.name}")
        return {
            "body_text": "", "tables": [], "pages": 0,
            "title": file_path.stem, "engine": "skipped(ocr_engine=none)",
            "extraction_errors": ["扫描件已跳过，如需处理请设置 OCR_ENGINE=auto 或 paddle/tesseract"],
        }

    if ocr_engine == "tesseract":
        return _extract_pdf_scanned_tesseract(file_path)

    if ocr_engine in ("paddle", "auto"):
        return _extract_pdf_scanned_paddleocr(file_path)   # 内部已有 auto 降级逻辑

    # 未知值降级
    logger.warning(f"[PDF] 未知 ocr_engine={ocr_engine!r}，降级为 auto")
    return _extract_pdf_scanned_paddleocr(file_path)


# ══════════════════════════════════════════════════════════════════════════════
# 其他格式解析器
# ══════════════════════════════════════════════════════════════════════════════
def _extract_docx(file_path: Path) -> dict:
    from docx import Document

    doc = Document(str(file_path))
    paragraphs: list = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading"):
            paragraphs.append(f"\n【{text}】")
        else:
            paragraphs.append(text)

    tables: list = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append({"rows": rows})
            '''
            姓名 | 年龄 | 城市
            张三 | 18 | 北京
            李四 | 20 | 上海
            '''
            paragraphs.append("\n[表格]\n" + "\n".join(" | ".join(r) for r in rows))

    props = doc.core_properties
    return {
        "body_text": "\n".join(paragraphs),
        "tables": tables,
        "title": props.title or file_path.stem,
        "author": props.author or "",
        "created_date": str(props.created) if props.created else "",
        "engine": "python-docx",
    }


def _extract_xlsx(file_path: Path) -> dict:
    import pandas as pd

    xl = pd.ExcelFile(str(file_path))
    sheets_text, tables = [], []        # sheets_text：每个sheet的文本内容，tables：每个sheet的表格数据
    for sheet_name in xl.sheet_names:   # xl.sheet_names：所有sheet的名称
        df = xl.parse(sheet_name)       # 解析当前sheet为DataFrame
        if df.empty:
            continue
        df = df.fillna("")              # 填充缺失值为空字符串
        tables.append({"sheet": sheet_name, "rows": df.values.tolist(), "columns": list(df.columns)})
        sheets_text.append(
            f"[Sheet: {sheet_name}]\n列名：{', '.join(str(c) for c in df.columns)}\n"   # 列名:姓名, 年龄, 城市
            f"{df.to_string(index=False, max_rows=500)}"
        )
    return {"body_text": "\n\n".join(sheets_text), "tables": tables,
            "title": file_path.stem, "engine": "pandas-xlsx"}


def _extract_csv(file_path: Path) -> dict:
    import pandas as pd

    df = None
    for enc in ("utf-8", "gbk", "utf-8-sig", "latin-1"):
        try:
            df = pd.read_csv(str(file_path), encoding=enc)
            break
        except Exception:
            continue
    if df is None:
        # 如果文件编码未知，强行 UTF-8 会报错，errors="replace" 可以让 pandas 跳过坏字符，这样至少能得到一个 DataFrame，而不是整个解析失败
        df = pd.read_csv(str(file_path), encoding="utf-8", errors="replace")
    body_text = f"列名：{', '.join(str(c) for c in df.columns)}\n{df.fillna('').to_string(index=False)}"
    return {"body_text": body_text,
            "tables": [{"rows": df.fillna("").values.tolist(), "columns": list(df.columns)}],
            "title": file_path.stem, "engine": "pandas-csv"}


def _extract_html(file_path: Path) -> dict:
    from bs4 import BeautifulSoup
    
    # errors="ignore" 用于忽略编码错误，避免解析失败,"html.parser" 是 BeautifulSoup 默认的解析器
    soup = BeautifulSoup(file_path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    title_tag = soup.find("title")
    title = title_tag.text.strip() if title_tag else file_path.stem   # file_path.stem：文件名（不包含扩展名）或文件路径
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe"]):
        tag.decompose()   # 移除这些标签，避免干扰解析
    tables: list = []
    for tbl in soup.find_all("table"):
        rows = [[td.get_text(strip=True) for td in tr.find_all(["td","th"])]
                for tr in tbl.find_all("tr")]
        rows = [r for r in rows if r]   # 过滤掉空行,[["姓名", "年龄"], ["张三", "18"]]
        if rows:
            tables.append({"rows": rows})
    return {"body_text": soup.get_text(separator="\n"), "title": title,
            "tables": tables, "engine": "beautifulsoup"}


def _extract_json(file_path: Path) -> dict:
    data = json.loads(file_path.read_text(encoding="utf-8"))
    body_text = ("\n".join(json.dumps(i, ensure_ascii=False) for i in data)     # 列表里的每个元素单独转成 JSON
                 if isinstance(data, list)
                 else json.dumps(data, ensure_ascii=False, indent=2))   # indent=2：缩进 2 个空格
    return {"body_text": body_text, "title": file_path.stem, "engine": "json"}


def _extract_text(file_path: Path) -> dict:
    return {"body_text": file_path.read_text(encoding="utf-8", errors="ignore"),
            "title": file_path.stem, "engine": "plain-text"}


# ══════════════════════════════════════════════════════════════════════════════
# 分发器
# ══════════════════════════════════════════════════════════════════════════════
_EXTRACTOR_MAP = {                           # 字典映射：文件类型 → 对应的解析函数。扩展新格式只需在这里加一行
    DocType.PDF:   _extract_pdf_enterprise,
    DocType.DOCX:  _extract_docx,
    DocType.XLSX:  _extract_xlsx,
    DocType.CSV:   _extract_csv,
    DocType.HTML:  _extract_html,
    DocType.JSON:  _extract_json,
    DocType.TXT:   _extract_text,
    DocType.MD:    _extract_text,
    DocType.IMAGE: None,   # standalone images bypass text extraction; handled separately in ExtractorService.extract()
}


def _detect_doc_type(path: Path) -> DocType:
    m = {".pdf": DocType.PDF, ".docx": DocType.DOCX, ".doc": DocType.DOCX,
         ".xlsx": DocType.XLSX, ".xls": DocType.XLSX, ".csv": DocType.CSV,
         ".html": DocType.HTML, ".htm": DocType.HTML, ".json": DocType.JSON,
         ".txt": DocType.TXT, ".md": DocType.MD,
         ".jpg": DocType.IMAGE, ".jpeg": DocType.IMAGE, ".png": DocType.IMAGE, ".webp": DocType.IMAGE}
    return m.get(path.suffix.lower(), DocType.UNKNOWN)


# ══════════════════════════════════════════════════════════════════════════════
# ExtractorService（对外接口不变）
# ══════════════════════════════════════════════════════════════════════════════
class ExtractorService:

    @log_latency
    async def extract(
        self,
        doc: RawDocument,
        llm_client=None,        # 传入时，扫描 PDF 可选用 Claude Vision 提取
    ) -> ExtractedContent:
        path = Path(doc.file_path)
        logger.info(f"[Extract] START raw_id={doc.raw_id} file={path.name}")    # START raw_id: 原始文档的唯一 ID

        if not path.exists():
            return ExtractedContent(
                raw_id=doc.raw_id, doc_type=doc.doc_type,
                extraction_errors=[f"File not found: {path}"],
            )

        doc_type = doc.doc_type if doc.doc_type != DocType.UNKNOWN else _detect_doc_type(path)
        extractor_fn = _EXTRACTOR_MAP.get(doc_type)     # extractor_fn：根据文件类型获取对应的解析函数

        if extractor_fn is None and doc_type != DocType.IMAGE:
            return ExtractedContent(
                raw_id=doc.raw_id, doc_type=doc_type,
                extraction_errors=[f"不支持的文件格式: {doc_type}"],
            )

        errors: list = []
        result_dict: dict = {}

        # ── Stage 2a: text extraction (skip for standalone image files) ─────────
        if doc_type != DocType.IMAGE and extractor_fn is not None:
            try:
                # Vision 路径：ocr_engine=vision + Anthropic llm_client + 扫描件 → 异步直接调用
                use_vision = (
                    doc_type == DocType.PDF
                    and getattr(settings, "ocr_engine", "auto") == "vision"
                    and llm_client is not None
                    and hasattr(llm_client, "chat_with_vision")
                    and is_scanned_pdf(path)
                )
                if use_vision:
                    logger.info(f"[Extract] 扫描PDF → Claude Vision: {path.name}")
                    result_dict = await _extract_pdf_vision_async(path, llm_client)
                else:
                    loop = asyncio.get_event_loop()
                    result_dict = await loop.run_in_executor(None, extractor_fn, path)

                if "extraction_errors" in result_dict:
                    errors.extend(result_dict["extraction_errors"])
            except Exception as exc:
                logger.error(f"[Extract] FAILED raw_id={doc.raw_id}: {exc}")
                errors.append(str(exc))

        # 提取完成后还要再清洗一遍，因为 OCR/Word 解析可能带入格式字符
        body_text = clean_text(result_dict.get("body_text", ""))
        content = ExtractedContent(
            raw_id=doc.raw_id,
            doc_type=doc_type,
            title=result_dict.get("title", path.stem),
            author=result_dict.get("author", ""),
            created_date=result_dict.get("created_date", ""),
            pages=result_dict.get("pages", 0),
            body_text=body_text,
            tables=result_dict.get("tables", []),
            images_count=result_dict.get("images_count", 0),
            metadata=result_dict.get("metadata", {}),
            extraction_engine=result_dict.get("engine", "unknown"),
            extraction_errors=errors,
        )

        # ── Stage 2b: image extraction (PDF only) ────────────────────────────
        if doc_type == DocType.PDF:
            loop2 = asyncio.get_event_loop()
            images = await loop2.run_in_executor(
                None, extract_images_from_pdf, path, doc.raw_id
            )
            content.images = images
            if content.images:
                content.images_count = len(content.images)

        # ── Stage 2c: standalone image file ──────────────────────────────────
        if doc_type == DocType.IMAGE:
            img_svc = get_image_extractor()
            try:
                standalone_img = img_svc.extract_standalone(path)
            except ValueError as exc:
                errors.append(str(exc))
                content.extraction_errors = errors
            else:
                content.images = [standalone_img]
                content.images_count = 1

        logger.info(
            f"[Extract] DONE raw_id={doc.raw_id} engine={content.extraction_engine} "
            f"chars={len(body_text)} pages={content.pages} tables={len(content.tables)}"
        )
        return content


def get_extractor() -> ExtractorService:
    return ExtractorService()
